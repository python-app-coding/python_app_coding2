# coding = utf8

import os
import docx
from docx import shared
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.shared import Cm

demo_file = 'temp_docx_demo.docx'

# 创建文档，并设置文档的字体格式（宋体、五号字、黑色）
demodoc = docx.Document()		# 创建空白文档
# set font
demodoc.styles['Normal'].font.name = '宋体'
demodoc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
demodoc.styles['Normal'].font.size = Pt(10.5)   	# 5号字体相当 抄 于 10.5pt
demodoc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

# 向文档对象添加标题（heading）
phh = demodoc.add_heading('辛弃疾', level=1)
phh.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
phh.runs[0].font.color.rgb = RGBColor(100,100,0)
phh.runs[0].font.name = '黑体'
phh.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 添加第一个段落（辛弃疾简介），设置字体为仿宋体，
# 设置汉字字体需要使用一个特定的设置方式（.element.rPr.rFonts.set(qn('w:eastAsia'), [字体名])），
# 同时，设置段落的对齐方式（JUSTIFY）和首行缩进量（Pt（21），按照每个字符10.5磅，为缩进两个字符）
# add first paragraph
ph0 = demodoc.add_paragraph()
ph0.paragraph_format.space_before = Pt(8)
run01 = ph0.add_run('原字坦夫，后改字幼安，中年后别号稼轩居士，山东济南府历城县'
                    '（今山东省济南市历城区遥墙镇四风闸村）人。'
                    '南宋官员、将领、文学家，豪放派词人，有“词中之龙”之称。'
                    '与苏轼合称“苏辛”，与李清照并称“济南二安”。')
run01.font.name = '仿宋'
run01.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run01.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
ph0.paragraph_format.first_line_indent = Pt(21)

# 添加词牌名词名标题（heading），对齐格式为居中（CENTER），字体设为楷体、红色，标题段落之后空间为10.5磅
# add heading
phh2 = demodoc.add_heading('最著名的词作\n破阵子·为陈同甫赋壮词以寄之', level=2)
phh2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
phh2.runs[0].font.color.rgb = RGBColor(255,10,10)
phh2.runs[0].font.name = '楷体'
phh2.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
phh2.paragraph_format.space_after = Pt(10.5)

# 添加词体上阙、下阙段落（paragraph）、区块（run），设置区块的字体为Cambria、加黑（bold），
# 设置区块的字体颜色（font.color.rgb）为OliveDrab4（105，139，34）：
# add paragraph
p1 = demodoc.add_paragraph()
r11 = p1.add_run('上阕: ')
r11.bold = True
r12 = p1.add_run('醉里挑灯看剑，梦回吹角连营。八百里分麾下炙，五十弦翻塞外声。沙场秋点兵。')
r12.font.color.rgb = RGBColor(105, 139, 34)
r12.font.name = 'Cambria'
r12.element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria')
r12.bold = True

p2 = demodoc.add_paragraph()
r21 = p2.add_run('下阙: ')
r21.bold = True
r22 = p2.add_run('马作的卢飞快，弓如霹雳弦惊。了却君王天下事，赢得生前身后名。可怜白发生！')
r22.font.color.rgb = RGBColor(30, 150, 100)
r22.font.name = '仿宋'
r22.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
r22.bold = True

# 添加一个段落，对齐方式为居中，在段落中增加一个区块，在该区块中添加辛弃疾的画像图片。
# 使用模块的__file__属性定位图片的路径（图片与模块在一个目录之中）。设置图片在文档中的宽和高分别为5厘米、6厘米
# add picture
# -- add paragraph, alignment to center
paragraph = demodoc.add_paragraph()
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = paragraph.add_run()
# -- set picture in run, set picture-path from __file__
file_path = os.path.split(os.path.abspath(__file__))[0]
run.add_picture(file_path + '/xin.png', width=shared.Cm(5), height=shared.Cm(6))

# 添加表格（风格为Light Shading），内容为辛弃疾编年表，设置表格的第二列对齐格式为居中，其余为缺省对齐格式（左对齐）。
# add table
tb = demodoc.add_table(0, 0, style='Light Shading')
# -- set table format：alignment
tb.alignment = WD_TABLE_ALIGNMENT.CENTER
# --增加表格为3列、8行，列宽度为3、2、5厘米，标题行高度为0.5厘米，其它各行高度为1厘米：
# --add columns, set columns width
tb.add_column(width=shared.Cm(3))
tb.add_column(width=shared.Cm(2))
tb.add_column(width=shared.Cm(5))
# -- add rows
row = tb.add_row()
row.height = docx.shared.Cm(0.5)
for _ in range(7):
    row = tb.add_row()
    row.height = docx.shared.Cm(1)
# --设置表格标题，添加表格数据
# -- set table header
table_header = ['年份', '年龄', '事项']
for ci, cell in enumerate(tb.rows[0].cells):
    cell.text = table_header[ci]
    cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell.paragraphs[0].runs[0].bold = True
# -- add table rows
content = [['宋绍兴十年（1140）', '出生', '生于山东历城之四风闸'],
           ['宋绍兴三十一年（1161）', '22岁', '金主亮大举南犯，稼轩聚众二千，与耿京共图恢复。'],
           ['宋绍兴三十二年（1162）', '23岁', '奉表南归。十八日至建康。召见，授右承务郎。'],
           ['宋淳熙十五年（1188）', '49岁', '友人陈同甫（亮）来访，长歌相答，作破阵子。'],
           ['宋庆元二年（1196）', '57岁', '以纠结徒党罪名罢斥朱熹及其门徒。'],
           ['宋嘉泰四年（1204）', '65岁', '支持韩侂胄发动对金战争。'],
           ['宋开禧三年（1207）', '68岁', '归铅山，八月得疾, 九月初十日卒。'],
           ]
for ri in range(len(content)):		                # 行遍历
    for ci in range(3):			                    # 列遍历
        tb.cell(ri+1, ci).text = content[ri][ci]    # 从第1行开始（跳过标题行）
    tb.cell(ri+1, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 250)
    tb.cell(ri+1, 1).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tb.cell(ri+1, 2).paragraphs[0].runs[0].font.name = '楷体'
    tb.cell(ri+1, 2).paragraphs[0].runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    tb.cell(ri + 1, 2).paragraphs[0].runs[0].bold = True

# 设置节区属性
# -- set section properties
section = demodoc.sections[0]
# 页边距设置
# -- set margins
section.top_margin = Cm(3.7)
section.bottom_margin = Cm(3.5)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.6)

# 保存文档对象为Word文件
demodoc.save(demo_file)
