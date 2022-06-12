# coding = utf8

import os
import docx


class DocxReader:
    """
    # 读出.docx文档的文本、表格、图片数据
    >>> dr = DocxReader('demo_docx.docx')		# 使用DocReader读取文档
    >>> doctext = dr.get_paragraphs()
    >>> for textline in doctext: print(textline)
    ['辛弃疾']
    ['原字坦夫，后改字幼安，中年后别号稼轩居士，山东济南府历城县（今山东省济南市历城区遥墙镇四风闸村）人。南宋官员、将领、文学家，豪放派词人，有“词中之龙”之称。与苏轼合称“苏辛”，与李清照并称“济南二安”。']
    ['最著名的词作\\n破阵子·为陈同甫赋壮词以寄之']
    ['上阕: ', '醉里挑灯看剑，梦回吹角连营。八百里分麾下炙，五十弦翻塞外声。沙场秋点兵。']
    ['下阙: ', '马作的卢飞快，弓如霹雳弦惊。了却君王天下事，赢得生前身后名。可怜白发生！']
    ['']

    >>> table_texts = dr.get_tables()
    >>> tbnum = len(table_texts)			        # 表格数
    >>> tbnum
    1
    >>> for textline in table_texts[0]:
    ...     print(textline)
    ['年份', '年龄', '事项']
    ['宋绍兴十年（1140）', '出生', '生于山东历城之四风闸']
    ['宋绍兴三十一年（1161）', '22岁', '金主亮大举南犯，稼轩聚众二千，与耿京共图恢复。']
    ['宋绍兴三十二年（1162）', '23岁', '奉表南归。十八日至建康。召见，授右承务郎。']
    ['宋淳熙十五年（1188）', '49岁', '友人陈同甫（亮）来访，长歌相答，作破阵子。']
    ['宋庆元二年（1196）', '57岁', '以纠结徒党罪名罢斥朱熹及其门徒。']
    ['宋嘉泰四年（1204）', '65岁', '支持韩侂胄发动对金战争。']
    ['宋开禧三年（1207）', '68岁', '归铅山，八月得疾, 九月初十日卒。']

    >>> pictures = dr.get_pictures('temp_path')
    >>> print(pictures)
    ['...\\temp_path\\xin.png']
    """

    def __init__(self, docxfile='demo.docx'):
        self.doc = docx.Document(docxfile)

    def get_paragraphs(self):
        """
        读取文档各段落的文本
        返回文档文本列表，元素为每段落各段程文本的列表
        """
        paragraphs = self.doc.paragraphs
        # 处理段落数据： 检索出段落中的文本，也可以进行其它处理，如搜索、修改等
        doctext = []
        for ph in paragraphs:
            text = []
            for run in ph.runs:
                text.append(run.text)
            doctext.append(text)
        return doctext

    def get_tables(self):
        """
        读取文档各表格的文本
        返回文档文本列表，元素为每表格各行各单元格文本列表的列表
        """
        tables = self.doc.tables
        tables_text = []
        for table in tables:
            text = []
            for row in table.rows:
                text.append([cell.text for cell in row.cells])
            tables_text.append(text)
        return tables_text

    def get_pictures(self, save_pictures_path='tempdir'):
        """
         读取文档中的图片数据
         返回写入临时目录中的图像文件名的列表
         """
        if not os.path.isdir(save_pictures_path):
            os.mkdir(save_pictures_path)
        save_pictures_path = os.path.abspath(save_pictures_path) + '\\'
        images = []
        for j in range(len(self.doc.inline_shapes)):
            rID = self.doc.inline_shapes[0]._inline.graphic.graphicData.pic.blipFill.blip.embed
            image = self.doc.part.related_parts[rID]
            # 处理图像数据：将图像数据写到文件，也可以选择进行其它处理
            image_file = save_pictures_path + 'image_{:03d}.png'.format(j + 1)
            with open(image_file, 'wb') as fp:
                fp.write(image._blob)
            images.append(image_file)
        return images


if __name__ == '__main__':
    pass
